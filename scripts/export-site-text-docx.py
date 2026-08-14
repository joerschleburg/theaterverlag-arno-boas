#!/usr/bin/env python3
"""Export all visible website texts from the built dist/ into a structured Word document."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
OUT = ROOT / "export" / "Theaterverlag_Arno_Boas_Website_Texte.docx"

SKIP_TAGS = {"script", "style", "noscript", "svg"}
SKIP_SELECTORS = [
    "#merk-root",
    ".cat-sticky",
    ".mobile-nav",
    "[data-mobile-nav]",
    ".toast",
    ".overlay",
]

MAIN_PAGE_ORDER = [
    ("index.html", "Startseite"),
    ("theaterstuecke/index.html", "Alle Theaterstücke"),
    ("autoren/index.html", "Autoren"),
    ("ueber-den-verlag/index.html", "Über den Verlag"),
    ("was-ist-ein-theaterverlag/index.html", "Was ist ein Theaterverlag?"),
    ("preise-und-bedingungen/index.html", "Preise & Bedingungen"),
    ("kontakt/index.html", "Kontakt"),
    ("impressum/index.html", "Impressum"),
    ("datenschutz/index.html", "Datenschutz"),
]


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text


def should_skip(el: Tag) -> bool:
    if el.name in SKIP_TAGS:
        return True
    for sel in SKIP_SELECTORS:
        try:
            if el.select_one(sel):
                return True
            if el.parent and el.parent.select_one(sel) and el in el.parent.select(sel):
                return True
        except Exception:
            pass
    if el.get("hidden") is not None:
        return True
    classes = " ".join(el.get("class", []))
    if "nav-toggle" in classes:
        return True
    return False


def iter_content_nodes(root: Tag):
    for el in root.descendants:
        if isinstance(el, NavigableString):
            continue
        if not isinstance(el, Tag) or should_skip(el):
            continue
        yield el


def page_path_to_label(path: Path) -> str:
  rel = path.relative_to(DIST).as_posix()
  if rel == "index.html":
    return "Startseite"
  parts = rel.replace("/index.html", "").split("/")
  return " / ".join(parts)


def extract_blocks(soup: BeautifulSoup, *, include_chrome: bool) -> list[tuple[str, str]]:
    body = soup.body
    if not body:
        return []

    blocks: list[tuple[str, str]] = []
    seen: set[str] = set()

    def add(kind: str, text: str):
        text = clean_text(text)
        if not text or text in seen:
            return
        seen.add(text)
        blocks.append((kind, text))

    # Meta
    title = clean_text(soup.title.string if soup.title else "")
    if title:
        add("meta_title", title)
    desc = soup.find("meta", attrs={"name": "description"})
    if desc and desc.get("content"):
        add("meta_description", desc["content"])

  # Remove chrome unless requested
    for sel in SKIP_SELECTORS:
        for node in body.select(sel):
            node.decompose()
    if not include_chrome:
        for node in body.select("header, footer, .site-mast, .ihead, nav[aria-label='Hauptnavigation'], nav[aria-label='Mobiles Menü']"):
            node.decompose()

    # Structured walk
    for el in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "dt", "dd", "blockquote", "label", "summary", "button", "a", "span", "div", "article"], recursive=True):
        if should_skip(el):
            continue
        if el.name in {"div", "article", "span"}:
            # only leaf-ish containers
            if el.find(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "details", "summary", "blockquote"]):
                continue
        text = clean_text(el.get_text(" ", strip=True))
        if not text:
            continue
        if el.name.startswith("h"):
            add(el.name, text)
        elif el.name == "summary":
            add("faq_q", text)
        elif el.parent and el.parent.name == "summary":
            continue
        elif el.name == "label":
            add("label", text)
        elif el.name == "blockquote":
            add("quote", text)
        elif el.name == "dt":
            add("term", text)
        elif el.name == "dd":
            add("definition", text)
        elif el.name == "li":
            add("list", text)
        elif el.name == "p":
            add("p", text)
        elif el.name in {"button", "a"} and len(text) < 120:
            add("cta", text)

    # FAQ answers in details
    for details in body.find_all("details"):
        summary = details.find("summary")
        answer = details.find("p")
        if summary and answer:
            q = clean_text(summary.get_text(" ", strip=True))
            a = clean_text(answer.get_text(" ", strip=True))
            if q and a:
                add("faq_q", q)
                add("faq_a", a)

    return blocks


def add_blocks_to_doc(doc: Document, blocks: list[tuple[str, str]]):
    for kind, text in blocks:
        if kind == "meta_title":
            continue
        if kind == "meta_description":
            p = doc.add_paragraph(text)
            p.runs[0].italic = True
            continue
        if kind.startswith("h"):
            level = int(kind[1])
            doc.add_heading(text, level=min(level, 4))
            continue
        if kind == "faq_q":
            p = doc.add_paragraph()
            r = p.add_run(f"Frage: {text}")
            r.bold = True
            continue
        if kind == "faq_a":
            doc.add_paragraph(f"Antwort: {text}")
            continue
        if kind == "quote":
            p = doc.add_paragraph(text)
            p.runs[0].italic = True
            continue
        if kind in {"term", "label", "cta"}:
            p = doc.add_paragraph()
            p.add_run(text).bold = True
            continue
        if kind == "definition":
            doc.add_paragraph(text)
            continue
        if kind == "list":
            doc.add_paragraph(text, style="List Bullet")
            continue
        doc.add_paragraph(text)


def extract_piece_list(soup: BeautifulSoup) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    for art in soup.select("article.li[data-titel]"):
        title = clean_text(art.get("data-titel", ""))
        subtitle = clean_text(art.get("data-untertitel", ""))
        autor = clean_text(art.get("data-autor", ""))
        inhalt = clean_text(art.get("data-inhalt", ""))
        teaser = inhalt[:220].rsplit(" ", 1)[0] + "…" if len(inhalt) > 220 else inhalt
        meta_parts = [
            subtitle,
            autor and f"Autor: {autor}",
            art.get("data-dauer", ""),
        ]
        meta = " · ".join(p for p in meta_parts if p)
        body = f"{meta}\n{teaser}" if teaser else meta
        if title:
            items.append((title, body))
    return sorted(items, key=lambda x: x[0].lower())


def collect_pages() -> dict[str, Path]:
    pages: dict[str, Path] = {}
    for path in DIST.rglob("index.html"):
        rel = path.relative_to(DIST).as_posix()
        pages[rel] = path
    if (DIST / "index.html").exists():
        pages["index.html"] = DIST / "index.html"
    return pages


def main():
    pages = collect_pages()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title page
    doc.add_heading("Theaterverlag Arno Boas", 0)
    doc.add_heading("Vollständiger Website-Textexport", level=1)
    doc.add_paragraph(f"Stand: {date.today().strftime('%d.%m.%Y')}")
    doc.add_paragraph(
        "Dieses Dokument enthält die sichtbaren Texte der Website in strukturierter Form — "
        "Hauptseiten, Autorenprofile und alle Theaterstück-Detailseiten."
    )
    doc.add_page_break()

    # Table of contents (manual)
    doc.add_heading("Inhaltsverzeichnis", level=1)
    toc = [label for _, label in MAIN_PAGE_ORDER]
    toc.append("Autorenprofile (Einzeln)")
    toc.append("Theaterstücke — Übersicht (alle Titel)")
    toc.append("Theaterstücke — Detailseiten")
    for i, item in enumerate(toc, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")
    doc.add_page_break()

    # Main pages
    doc.add_heading("Teil 1 — Hauptseiten", level=1)
    for rel, label in MAIN_PAGE_ORDER:
        path = pages.get(rel)
        if not path:
            continue
        soup = BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")
        doc.add_heading(label, level=2)
        doc.add_paragraph(f"Quelle: /{rel.replace('/index.html', '') or ''}")
        blocks = extract_blocks(soup, include_chrome=False)
        add_blocks_to_doc(doc, blocks)

        if rel == "theaterstuecke/index.html":
            doc.add_heading("Komplette Stückliste (Titel & Kurztext)", level=3)
            for title, body in extract_piece_list(soup):
                p = doc.add_paragraph()
                p.add_run(title).bold = True
                if body:
                    doc.add_paragraph(body)
        doc.add_page_break()

    # Authors
    author_pages = sorted(
        [p for rel, p in pages.items() if rel.startswith("autoren/") and rel != "autoren/index.html"],
        key=lambda p: p.parent.name,
    )
    if author_pages:
        doc.add_heading("Teil 2 — Autorenprofile", level=1)
        for path in author_pages:
            soup = BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")
            slug = path.parent.name
            doc.add_heading(slug.replace("-", " ").title(), level=2)
            doc.add_paragraph(f"Quelle: /autoren/{slug}")
            blocks = extract_blocks(soup, include_chrome=False)
            add_blocks_to_doc(doc, blocks)
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Piece detail pages
    piece_pages = sorted(
        [p for rel, p in pages.items() if rel.startswith("stuecke/")],
        key=lambda p: p.parent.name,
    )
    doc.add_heading("Teil 3 — Theaterstück-Detailseiten", level=1)
    for path in piece_pages:
        soup = BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")
        slug = path.parent.name
        title_tag = clean_text(soup.title.string.split("—")[0] if soup.title and soup.title.string else slug)
        doc.add_heading(title_tag, level=2)
        doc.add_paragraph(f"Quelle: /stuecke/{slug}")
        blocks = extract_blocks(soup, include_chrome=False)
        add_blocks_to_doc(doc, blocks)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Written: {OUT}")
    print(f"Pages exported: {len(MAIN_PAGE_ORDER)} main + {len(author_pages)} authors + {len(piece_pages)} pieces")


if __name__ == "__main__":
    main()

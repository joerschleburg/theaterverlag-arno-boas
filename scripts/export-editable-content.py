#!/usr/bin/env python3
"""Export import-friendly Word (static pages with IDs) and Excel (Stücke + Autoren)."""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
EXPORT = ROOT / "export"
MANIFEST_PATH = EXPORT / "content-manifest.json"
WORD_PATH = EXPORT / "Theaterverlag_Arno_Boas_Texte_Korrektur.docx"
EXCEL_PATH = EXPORT / "Theaterverlag_Arno_Boas_Stuecke_Korrektur.xlsx"
GUIDE_PATH = EXPORT / "KORREKTUR_ANLEITUNG.txt"

SITE = {
    "name": "Theaterverlag Arno Boas",
    "email": "info@theaterverlag-arno-boas.de",
    "phone": "07933/20093",
    "fax": "07933/20094",
    "founded": "1997",
    "address": {"street": "Finsterlohr 46", "zip": "97993", "city": "Creglingen"},
    "vatId": "DE263971681",
}

PAGE_LABELS = {
    "component.hero": "Startseite — Hero",
    "component.newsletter": "Newsletter (alle Seiten)",
    "index": "Startseite",
    "theaterstuecke": "Alle Theaterstücke",
    "autoren": "Autoren — Übersicht",
    "ueber-den-verlag": "Über den Verlag",
    "was-ist-ein-theaterverlag": "Was ist ein Theaterverlag?",
    "preise-und-bedingungen": "Preise & Bedingungen",
    "kontakt": "Kontakt",
    "impressum": "Impressum",
    "datenschutz": "Datenschutz",
}

ASTRO_PAGES = [
    ("pages/index.astro", "index"),
    ("pages/theaterstuecke/index.astro", "theaterstuecke"),
    ("pages/autoren/index.astro", "autoren"),
    ("pages/ueber-den-verlag.astro", "ueber-den-verlag"),
    ("pages/was-ist-ein-theaterverlag.astro", "was-ist-ein-theaterverlag"),
    ("pages/preise-und-bedingungen.astro", "preise-und-bedingungen"),
    ("pages/kontakt.astro", "kontakt"),
    ("pages/impressum.astro", "impressum"),
    ("pages/datenschutz.astro", "datenschutz"),
    ("components/Hero.astro", "component.hero"),
    ("components/NewsletterSignup.astro", "component.newsletter"),
]


def clean_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def resolve_site(text: str) -> str:
    text = text.replace("{SITE.name}", SITE["name"])
    text = text.replace("{SITE.email}", SITE["email"])
    text = text.replace("{SITE.phone}", SITE["phone"])
    text = text.replace("{SITE.fax}", SITE["fax"])
    text = text.replace("{SITE.founded}", SITE["founded"])
    text = text.replace("{SITE.vatId}", SITE["vatId"])
    text = text.replace("{SITE.address.street}", SITE["address"]["street"])
    text = text.replace("{SITE.address.zip}", SITE["address"]["zip"])
    text = text.replace("{SITE.address.city}", SITE["address"]["city"])
    text = re.sub(r"\$\{SITE\.email\}", SITE["email"], text)
    text = re.sub(r"\$\{SITE\.name\}", SITE["name"], text)
    text = re.sub(r"\$\{SITE\.founded\}", SITE["founded"], text)
    return clean_ws(text)


def split_astro(path: Path) -> tuple[str, str]:
    parts = path.read_text(encoding="utf-8").split("---", 2)
    if len(parts) < 3:
        return "", path.read_text(encoding="utf-8")
    return parts[1], parts[2]


def strip_astro_template(html: str) -> str:
    html = re.sub(r"<style[\s\S]*?</style>", "", html, flags=re.I)
    html = re.sub(r"<script[\s\S]*?</script>", "", html, flags=re.I)
    html = re.sub(r"\{[\s\S]*?\}", "", html)
    html = re.sub(r"\[\s*\.\.\.[^\]]*\]", "", html)
    return html


def extract_faqs(content: str, page_id: str, rel_file: str) -> list[dict]:
    entries: list[dict] = []
    block = re.search(r"const faqs\s*=\s*\[(.*?)\];", content, re.S)
    if not block:
        return entries
    body = block.group(1)
    chunks = re.split(r"\},\s*\{", body)
    for i, chunk in enumerate(chunks):
        qm = re.search(r'question:\s*(?:"((?:\\.|[^"\\])*)"|`((?:\\.|[^`\\])*)`)', chunk, re.S)
        am = re.search(r'answer:\s*(?:"((?:\\.|[^"\\])*)"|`((?:\\.|[^`\\])*)`)', chunk, re.S)
        if not qm or not am:
            continue
        q = qm.group(1) or qm.group(2) or ""
        a = am.group(1) or am.group(2) or ""
        q = bytes(q, "utf-8").decode("unicode_escape") if "\\" in q else q
        a = bytes(a, "utf-8").decode("unicode_escape") if "\\" in a else a
        q = resolve_site(q.replace('\\"', '"'))
        a = resolve_site(a.replace('\\"', '"'))
        for field, text in [("question", q), ("answer", a)]:
            eid = f"{page_id}.faq.{i}.{field}"
            entries.append(
                {
                    "id": eid,
                    "section": PAGE_LABELS.get(page_id, page_id),
                    "label": f"FAQ {i + 1} — {'Frage' if field == 'question' else 'Antwort'}",
                    "text": text,
                    "import": {
                        "file": rel_file,
                        "kind": "faq",
                        "index": i,
                        "field": field,
                    },
                }
            )
    return entries


def extract_index_frontmatter(fm: str) -> list[dict]:
    entries: list[dict] = []
    rel = "src/pages/index.astro"

    def add(eid: str, label: str, text: str):
        if not clean_ws(text):
            return
        entries.append(
            {
                "id": f"index.{eid}",
                "section": PAGE_LABELS["index"],
                "label": label,
                "text": resolve_site(text),
                "import": {
                    "file": rel,
                    "kind": "replace",
                    "original": text if isinstance(text, str) else resolve_site(text),
                },
            }
        )

    for key, label in [
        ("kicker", "Im Rampenlicht — Kicker"),
        ("empfehlung", "Im Rampenlicht — Empfehlung"),
        ("headline", "Im Rampenlicht — Überschrift"),
        ("dek", "Im Rampenlicht — Teaser"),
    ]:
        m = re.search(rf"{key}:\s*['\"](.*?)['\"]", fm, re.S)
        if m:
            add(f"feature.{key}", label, m.group(1))

    titel = re.search(r"titelLines:\s*\[(.*?)\]", fm, re.S)
    if titel:
        lines = re.findall(r"['\"](.*?)['\"]", titel.group(1))
        if lines:
            add("feature.titel", "Im Rampenlicht — Titel", " / ".join(lines))

    neu_blocks = re.findall(
        r"\{\s*href:[\s\S]*?titel:\s*['\"](.*?)['\"][\s\S]*?\}",
        fm,
    )
    for i, titel in enumerate(neu_blocks):
        add(f"neu.{i}.titel", f"Neu im Repertoire — Titel {i + 1}", titel)

    return entries


def extract_friends(fm: str) -> list[dict]:
    entries: list[dict] = []
    rel = "src/pages/ueber-den-verlag.astro"
    blocks = re.findall(
        r"name:\s*\"([^\"]+)\",\s*desc:\s*\"([^\"]+)\"",
        fm,
    )
    for i, (name, desc) in enumerate(blocks):
        eid = f"ueber-den-verlag.friends.{i}.desc"
        entries.append(
            {
                "id": eid,
                "section": PAGE_LABELS["ueber-den-verlag"],
                "label": f"Verlagsfreunde — {name}",
                "text": desc,
                "import": {
                    "file": rel,
                    "kind": "replace",
                    "original": desc,
                },
            }
        )
    return entries


def extract_template_texts(html: str, page_id: str, rel_file: str) -> list[dict]:
    entries: list[dict] = []
    counters: dict[str, int] = {}

    def bump(kind: str) -> int:
        counters[kind] = counters.get(kind, 0) + 1
        return counters[kind]

    soup = BeautifulSoup(strip_astro_template(html), "lxml")

    # FAQs werden separat extrahiert
    for el in soup.select(".faq summary, .faq p"):
        el.decompose()

    selectors = [
        ("h1", "h1"),
        ("h2", "h2"),
        ("h3", "h3"),
        ("p.lead", "lead"),
        ("p.big", "big"),
        ("p.dek", "dek"),
        ("span.kick", "kick"),
        ("blockquote", "quote"),
        ("p", "p"),
        ("li", "li"),
        ("dt", "dt"),
        ("dd", "dd"),
        ("span.step-title", "step_title"),
        ("span.step-text", "step_text"),
        ("label", "label"),
        ("summary", "summary"),
        ("div.hint", "hint"),
        ("div.nl-note", "note"),
    ]

    seen: set[str] = set()
    for sel, kind in selectors:
        for el in soup.select(sel):
            text = clean_ws(el.get_text(" ", strip=True))
            if not text or len(text) < 2:
                continue
            if text in seen:
                continue
            seen.add(text)
            idx = bump(kind)
            eid = f"{page_id}.{kind}.{idx}"
            entries.append(
                {
                    "id": eid,
                    "section": PAGE_LABELS.get(page_id, page_id),
                    "label": f"{kind.replace('_', ' ').title()} {idx}",
                    "text": resolve_site(text),
                    "import": {
                        "file": rel_file,
                        "kind": "replace",
                        "original": text,
                    },
                }
            )
    return entries


def parse_autoren() -> list[dict]:
    text = (SRC / "lib" / "autoren.ts").read_text(encoding="utf-8")
    items: list[dict] = []
    for block in re.findall(r"\{[^{}]*slug:\s*\"([^\"]+\"[^{}]*)\}", text):
        pass
    # multiline blocks
    for m in re.finditer(
        r"\{\s*slug:\s*\"([^\"]+)\",\s*name:\s*\"([^\"]+)\",\s*role:\s*\"([^\"]+)\"(?:,\s*photo:[\s\S]*?)?(?:,\s*photoStock:[\s\S]*?)?(?:,\s*bio:\s*\"([^\"]*)\")?",
        text,
    ):
        items.append(
            {
                "slug": m.group(1),
                "name": m.group(2),
                "role": m.group(3),
                "bio": m.group(4) or "",
            }
        )
    return items


def build_manifest() -> dict:
    entries: list[dict] = []
    for rel, page_id in ASTRO_PAGES:
        path = SRC / rel
        if not path.exists():
            continue
        full = path.read_text(encoding="utf-8")
        rel_file = f"src/{rel}"
        entries.extend(extract_faqs(full, page_id, rel_file))
        if page_id == "index":
            fm, html = split_astro(path)
            entries.extend(extract_index_frontmatter(fm))
        elif page_id == "ueber-den-verlag":
            fm, html = split_astro(path)
            entries.extend(extract_friends(fm))
        else:
            _, html = split_astro(path)
        entries.extend(extract_template_texts(html, page_id, rel_file))

    # dedupe by id
    by_id: dict[str, dict] = {}
    for e in entries:
        by_id[e["id"]] = e
    return {
        "generated": date.today().isoformat(),
        "entries": list(by_id.values()),
    }


def add_id_paragraph(doc: Document, entry_id: str):
    p = doc.add_paragraph()
    run = p.add_run(f"[ID: {entry_id}]")
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.size = Pt(9)
    run.italic = True


def write_word(manifest: dict):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("Theaterverlag Arno Boas — Textkorrektur", 0)
    doc.add_paragraph(f"Stand: {date.today().strftime('%d.%m.%Y')}")
    doc.add_paragraph(
        "Bitte nur den Text UNTER jeder grauen ID-Zeile korrigieren. "
        "Die Zeilen mit [ID: …] nicht ändern, löschen oder verschieben."
    )
    doc.add_page_break()

    current_section = None
    for entry in sorted(manifest["entries"], key=lambda e: (e["section"], e["id"])):
        if entry["section"] != current_section:
            current_section = entry["section"]
            doc.add_heading(current_section, level=1)
        doc.add_paragraph(entry["label"]).runs[0].bold = True
        add_id_paragraph(doc, entry["id"])
        doc.add_paragraph(entry["text"])

    doc.save(WORD_PATH)


def write_excel(stuecke: list[dict], autoren: list[dict]):
    wb = Workbook()

    # Sheet Anleitung
    ws0 = wb.active
    ws0.title = "Anleitung"
    lines = [
        "Korrektur-Anleitung — Theaterverlag Arno Boas",
        "",
        "STÜCKE-Blatt:",
        "• Spalte slug und titel NICHT ändern (technische Schlüssel).",
        "• Nur untertitel, inhalt und buehnenbild bearbeiten.",
        "• Keine Zeilen löschen oder neue einfügen.",
        "",
        "AUTOREN-Blatt:",
        "• Spalte slug NICHT ändern.",
        "• name, role und bio dürfen korrigiert werden.",
        "",
        "WORD-Datei (Texte_Korrektur.docx):",
        "• Nur Text unter [ID: …]-Zeilen ändern, IDs selbst unverändert lassen.",
    ]
    for i, line in enumerate(lines, 1):
        ws0.cell(row=i, column=1, value=line)
    ws0.column_dimensions["A"].width = 90

    header_fill = PatternFill("solid", fgColor="D9D9D9")
    locked_fill = PatternFill("solid", fgColor="F2F2F2")

    def setup_sheet(ws, headers, rows, locked_cols):
        ws.append(headers)
        for col in range(1, len(headers) + 1):
            c = ws.cell(row=1, column=col)
            c.font = Font(bold=True)
            c.fill = header_fill
            c.alignment = Alignment(wrap_text=True, vertical="top")
        for row in rows:
            ws.append(row)
        for i, width in enumerate([22, 36, 28, 70, 50], 1):
            ws.column_dimensions[get_column_letter(i)].width = width
        for r in range(2, ws.max_row + 1):
            for c in locked_cols:
                ws.cell(row=r, column=c).fill = locked_fill
            for c in range(1, len(headers) + 1):
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")

    ws1 = wb.create_sheet("Stuecke")
    stueck_rows = [
        [
            s.get("slug", ""),
            s.get("titel", ""),
            s.get("untertitel", ""),
            s.get("inhalt", ""),
            s.get("buehnenbild", ""),
        ]
        for s in sorted(stuecke, key=lambda x: (x.get("titel") or "").lower())
    ]
    setup_sheet(
        ws1,
        ["slug", "titel", "untertitel", "inhalt", "buehnenbild"],
        stueck_rows,
        locked_cols=[1, 2],
    )

    ws2 = wb.create_sheet("Autoren")
    autor_rows = [
        [a["slug"], a["name"], a["role"], a.get("bio", "")]
        for a in sorted(autoren, key=lambda x: x["name"].lower())
    ]
    setup_sheet(ws2, ["slug", "name", "role", "bio"], autor_rows, locked_cols=[1])

    wb.save(EXCEL_PATH)


def write_guide(manifest: dict, stuecke_count: int, autoren_count: int):
    GUIDE_PATH.write_text(
        "\n".join(
            [
                "Korrektur-Anleitung — Theaterverlag Arno Boas",
                f"Stand: {date.today().strftime('%d.%m.%Y')}",
                "",
                "Dateien:",
                f"1) {WORD_PATH.name} — Website-Texte mit festen IDs ({len(manifest['entries'])} Textblöcke)",
                f"2) {EXCEL_PATH.name} — {stuecke_count} Stücke + {autoren_count} Autoren",
                f"3) {MANIFEST_PATH.name} — technische Zuordnung für den Rückimport",
                "",
                "Word:",
                "- Nur den Text UNTER [ID: …] ändern",
                "- ID-Zeilen nicht anfassen",
                "",
                "Excel — Blatt Stuecke:",
                "- slug + titel nicht ändern",
                "- untertitel, inhalt, buehnenbild korrigieren",
                "",
                "Excel — Blatt Autoren:",
                "- slug nicht ändern",
                "- name, role, bio korrigieren",
                "",
                "Nach der Korrektur beide Dateien zurückschicken — dann können die Änderungen automatisch übernommen werden.",
            ]
        ),
        encoding="utf-8",
    )


def main():
    EXPORT.mkdir(parents=True, exist_ok=True)
    manifest = build_manifest()
    MANIFEST_PATH.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    stuecke = json.loads((SRC / "data" / "stuecke.json").read_text(encoding="utf-8"))[
        "stuecke"
    ]
    autoren = parse_autoren()

    write_word(manifest)
    write_excel(stuecke, autoren)
    write_guide(manifest, len(stuecke), len(autoren))

    print(f"Manifest: {MANIFEST_PATH} ({len(manifest['entries'])} entries)")
    print(f"Word:     {WORD_PATH}")
    print(f"Excel:    {EXCEL_PATH}")
    print(f"Guide:    {GUIDE_PATH}")


if __name__ == "__main__":
    main()
